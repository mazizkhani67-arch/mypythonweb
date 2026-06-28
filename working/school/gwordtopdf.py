import pandas as pd
from docx import Document
from docx2pdf import convert
import os
from datetime import datetime
import re

def fill_word_template(template_path, data, output_path):
    """
    Fill a Word template with data using various placeholder formats
    """
    # Load the template
    doc = Document(template_path)
    
    print(f"\n📝 Filling template: {os.path.basename(template_path)}")
    print(f"   Data: {data}")
    
    # Method 1: Replace in paragraphs
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        new_text = original_text
        
        # Try different placeholder formats
        for key, value in data.items():
            # Format 1: {{key}}
            placeholder1 = f'{{{{{key}}}}}'
            if placeholder1 in new_text:
                new_text = new_text.replace(placeholder1, str(value))
                print(f"   ✅ Replaced {placeholder1} → {value}")
            
            
            
           
        
        # Update the paragraph text if changed
        if new_text != original_text:
            paragraph.text = new_text
    
    # Method 2: Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    new_text = original_text
                    
                    for key, value in data.items():
                        # Try all placeholder formats
                        placeholders = [
                            f'{{{{{key}}}}}',  # {{key}}
                            f'{{{key}}}',       # {key}
                            f'<<{key}>>',       # <<key>>
                            f'${key}$',         # $key$
                            f'[{key}]',         # [key]
                            f'<{key}>'          # <key>
                        ]
                        
                        for placeholder in placeholders:
                            if placeholder in new_text:
                                new_text = new_text.replace(placeholder, str(value))
                                print(f"   ✅ Replaced {placeholder} → {value}")
                    
                    if new_text != original_text:
                        paragraph.text = new_text
    
    # Save the filled document
    doc.save(output_path)
    print(f"   💾 Saved: {os.path.basename(output_path)}")
    return output_path

def process_excel_with_templates(excel_file, template_folder, output_folder='pdf_output'):
    """
    Process Excel data using group-specific templates
    """
    # Read Excel
    df = pd.read_excel(excel_file)
    
    print("=" * 60)
    print("📊 EXCEL DATA")
    print("=" * 60)
    print(f"Columns: {df.columns.tolist()}")
    print(f"Total records: {len(df)}")
    print("\nFirst 2 rows:")
    print(df.head(2))
    
    # Map column names (handle both uppercase and lowercase)
    column_map = {}
    for col in df.columns:
        col_lower = col.lower().strip()
        if col_lower in ['name', 'firstname', 'first name', 'نام']:
            column_map['name'] = col
        elif col_lower in ['lastname', 'last name', 'surname', 'نام خانوادگی', 'family']:
            column_map['lastname'] = col
        elif col_lower in ['ncode', 'nationalcode', 'national code', 'code', 'id', 'کد ملی']:
            column_map['ncode'] = col
        elif col_lower in ['group', 'groups', 'class', 'section', 'گروه']:
            column_map['group'] = col
    
    # If no columns found with auto-detection, use exact column names
    if not column_map:
        print("\n⚠️ Auto-detection failed. Using exact column names...")
        for col in df.columns:
            if col in ['Name', 'name', 'FirstName', 'firstname']:
                column_map['name'] = col
            elif col in ['LastName', 'lastname', 'Surname', 'surname']:
                column_map['lastname'] = col
            elif col in ['NCode', 'ncode', 'NationalCode', 'nationalcode']:
                column_map['ncode'] = col
            elif col in ['Group', 'group', 'Groups', 'groups']:
                column_map['group'] = col
    
    # Check if we found all required columns
    required = ['name', 'lastname', 'ncode', 'group']
    missing = [r for r in required if r not in column_map]
    
    if missing:
        print(f"\n❌ Missing columns: {missing}")
        print(f"Available columns: {df.columns.tolist()}")
        print("\nPlease rename your columns to match:")
        print("  - 'name' or 'Name'")
        print("  - 'lastname' or 'LastName'")
        print("  - 'ncode' or 'NCode'")
        print("  - 'group' or 'Group'")
        return
    
    # Rename columns
    df = df.rename(columns={v: k for k, v in column_map.items()})
    print(f"\n✅ Column mapping: {column_map}")
    
    # Create output folder
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # Get unique groups
    groups = df['group'].unique()
    print(f"\n📁 Groups: {list(groups)}")
    
    # Check templates
    print("\n" + "=" * 60)
    print("🔍 CHECKING TEMPLATES")
    print("=" * 60)
    
    templates_found = []
    
    for group in groups:
        # Try different template name formats
        possible_names = [
            f"template_{group}.docx",
            f"template_{str(group).lower()}.docx",
            f"template_{str(group).upper()}.docx",
            f"group_{group}.docx",
            f"{group}.docx",
            f"Template_{group}.docx"
        ]
        
        template_path = None
        for name in possible_names:
            full_path = os.path.join(template_folder, name)
            if os.path.exists(full_path):
                template_path = full_path
                break
        
        if template_path:
            templates_found.append((group, template_path))
            print(f"✅ Group '{group}': {os.path.basename(template_path)}")
        else:
            print(f"❌ Group '{group}': No template found!")
            print(f"   Looking for: template_{group}.docx")
    
    if not templates_found:
        print("\n❌ No templates found! Please create template files.")
        print(f"Place them in: {template_folder}")
        print("Example: template_Group1.docx, template_Group2.docx, etc.")
        return
    
    # Process each group
    print("\n" + "=" * 60)
    print("🚀 GENERATING DOCUMENTS")
    print("=" * 60)
    
    total_created = 0
    
    for group, template_path in templates_found:
        group_folder = os.path.join(output_folder, f'group_{group}')
        if not os.path.exists(group_folder):
            os.makedirs(group_folder)
        
        group_data = df[df['group'] == group]
        print(f"\n👥 Group '{group}': {len(group_data)} people")
        
        for index, row in group_data.iterrows():
            try:
                # Prepare data
                data = {
                    'name': str(row['name']),
                    'lastname': str(row['lastname']),
                    'ncode': str(row['ncode']),
                    'group': str(row['group']),
                    'date': datetime.now().strftime('%Y-%m-%d'),
                    'datetime': datetime.now().strftime('%Y-%m-%d %H:%M')
                }
                
                # Generate filename
                filename = f"{row['name']}_{row['lastname']}_{row['ncode']}"
                filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_')).strip()
                filename = filename.replace(' ', '_')
                
                # Create Word document
                docx_path = os.path.join(group_folder, f"{filename}.docx")
                fill_word_template(template_path, data, docx_path)
                
                # Try to convert to PDF
                try:
                    pdf_path = os.path.join(group_folder, f"{filename}.pdf")
                    convert(docx_path, pdf_path)
                    print(f"  ✅ {filename}.pdf")
                    # Remove docx if PDF was created
                    # os.remove(docx_path)
                except Exception as e:
                    print(f"  ✅ {filename}.docx (PDF not created: {e})")
                
                total_created += 1
                
            except Exception as e:
                print(f"  ❌ Error: {e}")
    
    print("\n" + "=" * 60)
    print(f"✅ Created {total_created} documents")
    print(f"📁 Output: {output_folder}")
    print("=" * 60)

# Main execution
if __name__ == "__main__":
    # ============================================
    # UPDATE THESE PATHS
    # ============================================
    excel_file = "studentsDetails.xlsx"
    template_folder = "templates"
    # ============================================
    
    print("=" * 60)
    print("📄 DOCUMENT GENERATOR")
    print("=" * 60)
    
    # Check if files exist
    if not os.path.exists(excel_file):
        print(f"❌ Excel not found: {excel_file}")
        print(f"\nFiles in current directory:")
        for file in os.listdir('.'):
            print(f"  - {file}")
        exit()
    
    if not os.path.exists(template_folder):
        print(f"❌ Template folder not found: {template_folder}")
        print(f"\nCreating folder...")
        os.makedirs(template_folder)
        print(f"✅ Created: {template_folder}")
        print(f"\nPlease add your templates:")
        print("  - template_Group1.docx")
        print("  - template_Group2.docx")
        print("  - template_Group3.docx")
        print("  - template_Group4.docx")
        exit()
    
    # Show templates
    print("\n📄 Available templates:")
    templates = [f for f in os.listdir(template_folder) if f.endswith('.docx')]
    if templates:
        for t in templates:
            print(f"  - {t}")
    else:
        print("  ⚠️ No templates found!")
        exit()
    
    # Run the main process
    process_excel_with_templates(excel_file, template_folder, 'pdf_output')
    
    print("\n✅ Process completed!")